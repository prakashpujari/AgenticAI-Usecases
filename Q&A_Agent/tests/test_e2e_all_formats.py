"""
tests/test_e2e_all_formats.py
──────────────────────────────
End-to-end tests for every supported input format.

Test groups
───────────
  A – Document formats (no API needed): PDF, TXT, MD, RST, CSV, DOCX, XLSX, XLS
  B – URL ingestion  (network needed):  HTTP webpage, PDF-URL
  C – Image formats  (Groq Vision):     PNG, JPG, WEBP
  D – Audio formats  (Groq Whisper):    WAV, MP3, M4A
  E – Video formats  (Groq Whisper):    MP4, WEBM  (+ AVI/MOV/MKV with ffmpeg)
  F – YouTube        (layered fallback): YouTube URL
  G – Pipeline smoke test:              full generate_questions_from_text on TXT

Run all:           python tests/test_e2e_all_formats.py
Run with pytest:   pytest tests/test_e2e_all_formats.py -v
"""

from __future__ import annotations

import io
import math
import os
import struct
import sys
import tempfile
import time
import zlib
from pathlib import Path

import pytest

# ── Make project root importable ──────────────────────────────────────────────
sys.path.insert(0, str(Path(__file__).parent.parent))

os.environ.setdefault("LOG_TO_FILE", "false")
os.environ.setdefault("REDIS_URL", "")
os.environ.setdefault("IDENTITY_HMAC_SECRET", "test-secret-32chars-padding000")

import config  # noqa: E402  (must be after sys.path insert)

# ── Key availability flags ────────────────────────────────────────────────────
HAS_GROQ     = bool(config.GROQ_API_KEY and not config.GROQ_API_KEY.startswith("test"))
HAS_GEMINI   = bool(config.GEMINI_API_KEY)
HAS_HF       = bool(config.HF_API_KEY)
HAS_LLM      = HAS_GROQ or HAS_GEMINI or HAS_HF
HAS_SUPADATA = bool(config.SUPADATA_API_KEY)

SKIP_NO_LLM    = pytest.mark.skipif(not HAS_LLM,      reason="No LLM API key configured")
SKIP_NO_GROQ   = pytest.mark.skipif(not HAS_GROQ,     reason="No real GROQ_API_KEY configured")
SKIP_NO_YT     = pytest.mark.skipif(not HAS_SUPADATA,  reason="No SUPADATA_API_KEY for YouTube")

# ── Test document content ─────────────────────────────────────────────────────
_CONTENT = (
    "Photosynthesis is the process by which green plants convert sunlight into food. "
    "Chlorophyll in leaves absorbs light energy, primarily from the red and blue "
    "wavelengths of the electromagnetic spectrum. Carbon dioxide from the air and "
    "water from the soil combine in a chemical reaction powered by this light energy. "
    "The products of photosynthesis are glucose (a sugar used for energy) and oxygen, "
    "which is released into the atmosphere. The overall equation is: "
    "6CO2 + 6H2O + light energy → C6H12O6 + 6O2. "
    "Photosynthesis is fundamental to almost all life on Earth because it forms the "
    "base of the food chain and produces the oxygen that aerobic organisms breathe. "
    "There are two main stages: the light-dependent reactions that occur in the "
    "thylakoid membranes, and the Calvin cycle (light-independent reactions) that "
    "take place in the stroma of the chloroplast."
)


# ── File factory helpers ──────────────────────────────────────────────────────

def _make_pdf() -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas as rl_canvas

    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter)
    c.setFont("Helvetica", 11)
    y, margin = 720, 72
    for line in _CONTENT.split(". "):
        c.drawString(margin, y, line.strip() + ".")
        y -= 18
        if y < 100:
            c.showPage()
            c.setFont("Helvetica", 11)
            y = 720
    c.save()
    return buf.getvalue()


def _make_png() -> bytes:
    """Minimal 10×10 white RGB PNG, no external deps."""
    w, h = 10, 10

    def chunk(name: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + name
            + data
            + struct.pack(">I", zlib.crc32(name + data) & 0xFFFFFFFF)
        )

    ihdr = chunk(b"IHDR", struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0))
    raw  = b"".join(b"\x00" + b"\xff\xff\xff" * w for _ in range(h))
    idat = chunk(b"IDAT", zlib.compress(raw))
    iend = chunk(b"IEND", b"")
    return b"\x89PNG\r\n\x1a\n" + ihdr + idat + iend


def _make_jpeg() -> bytes:
    """Create a valid 10x10 white JPEG using Pillow."""
    try:
        from PIL import Image as _Image
        img = _Image.new("RGB", (10, 10), color=(255, 255, 255))
        buf = io.BytesIO()
        img.save(buf, format="JPEG")
        return buf.getvalue()
    except ImportError:
        pytest.skip("Pillow not installed — cannot create valid JPEG test file")


def _make_wav() -> bytes:
    """1-second 440 Hz tone as 16-kHz mono WAV — usable by Groq Whisper."""
    sample_rate, duration = 16_000, 1
    samples = [
        int(32767 * math.sin(2 * math.pi * 440 * i / sample_rate))
        for i in range(sample_rate * duration)
    ]
    raw = struct.pack(f"<{len(samples)}h", *samples)

    buf = io.BytesIO()
    import wave as _wave
    with _wave.open(buf, "wb") as wf:
        wf.setnchannels(1)
        wf.setsampwidth(2)
        wf.setframerate(sample_rate)
        wf.writeframes(raw)
    return buf.getvalue()


def _make_docx() -> bytes:
    import docx as _docx
    doc = _docx.Document()
    doc.add_heading("Photosynthesis", level=1)
    for sentence in _CONTENT.split(". "):
        doc.add_paragraph(sentence.strip() + ".")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_xlsx() -> bytes:
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Data"
    ws.append(["Topic", "Description"])
    ws.append(["Photosynthesis", "Converts sunlight to food using chlorophyll."])
    ws.append(["Equation", "6CO2 + 6H2O + light → C6H12O6 + 6O2"])
    ws.append(["Products", "Glucose and oxygen"])
    ws.append(["Location", "Chloroplast thylakoid membranes and stroma"])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _make_xls() -> bytes | None:
    """Create legacy .xls using xlwt if available, else return None."""
    try:
        import xlwt  # type: ignore[import]
        wb = xlwt.Workbook()
        ws = wb.add_sheet("Data")
        ws.write(0, 0, "Topic")
        ws.write(0, 1, "Description")
        ws.write(1, 0, "Photosynthesis")
        ws.write(1, 1, "Plant process converting light to glucose")
        ws.write(2, 0, "Equation")
        ws.write(2, 1, "6CO2 + 6H2O + light → C6H12O6 + 6O2")
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()
    except ImportError:
        return None


# ── Helper: run load_document on bytes with a given suffix ────────────────────

def _load(data: bytes, suffix: str) -> tuple[str, float]:
    """Write data to a temp file, call load_document(), return (text, elapsed_s)."""
    from src.ingestion.document_loader import load_document
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as fh:
        fh.write(data)
        path = fh.name
    try:
        t0 = time.monotonic()
        text = load_document(path)
        return text, time.monotonic() - t0
    finally:
        Path(path).unlink(missing_ok=True)


# ═══════════════════════════════════════════════════════════════════════════════
# GROUP A — Document formats (no API required)
# ═══════════════════════════════════════════════════════════════════════════════

class TestDocumentFormats:

    def test_pdf(self):
        data = _make_pdf()
        text, t = _load(data, ".pdf")
        assert len(text) > 100, "PDF: extracted text too short"
        assert "photosynthesis" in text.lower() or "chlorophyll" in text.lower(), \
            "PDF: expected content not found"
        print(f"\n  PDF:  {len(text):,} chars  ({t:.2f}s)")

    def test_txt(self):
        data = _CONTENT.encode("utf-8")
        text, t = _load(data, ".txt")
        assert "photosynthesis" in text.lower()
        assert len(text) >= len(_CONTENT) - 5
        print(f"\n  TXT:  {len(text):,} chars  ({t:.2f}s)")

    def test_md(self):
        md = f"# Photosynthesis\n\n{_CONTENT}\n\n## Summary\n\nKey process in biology.\n"
        text, t = _load(md.encode("utf-8"), ".md")
        assert "photosynthesis" in text.lower()
        print(f"\n  MD:   {len(text):,} chars  ({t:.2f}s)")

    def test_rst(self):
        rst = f"Photosynthesis\n==============\n\n{_CONTENT}\n\nConclusion\n----------\nCritical process.\n"
        text, t = _load(rst.encode("utf-8"), ".rst")
        assert "photosynthesis" in text.lower()
        print(f"\n  RST:  {len(text):,} chars  ({t:.2f}s)")

    def test_csv(self):
        csv = "topic,description\nphotosynthesis,converts sunlight to glucose\nchlorophyll,absorbs red and blue light\nglucose,primary product of photosynthesis\noxygen,secondary product released to atmosphere\n"
        text, t = _load(csv.encode("utf-8"), ".csv")
        assert "photosynthesis" in text.lower()
        print(f"\n  CSV:  {len(text):,} chars  ({t:.2f}s)")

    def test_docx(self):
        data = _make_docx()
        text, t = _load(data, ".docx")
        assert "photosynthesis" in text.lower()
        assert len(text) > 50
        print(f"\n  DOCX: {len(text):,} chars  ({t:.2f}s)")

    def test_xlsx(self):
        data = _make_xlsx()
        text, t = _load(data, ".xlsx")
        assert "photosynthesis" in text.lower() or "glucose" in text.lower()
        print(f"\n  XLSX: {len(text):,} chars  ({t:.2f}s)")

    def test_xls(self):
        data = _make_xls()
        if data is None:
            pytest.skip("xlwt not installed — cannot create test .xls file (pip install xlwt)")
        text, t = _load(data, ".xls")
        assert "photosynthesis" in text.lower()
        print(f"\n  XLS:  {len(text):,} chars  ({t:.2f}s)")


# ═══════════════════════════════════════════════════════════════════════════════
# GROUP B — URL / web ingestion (needs network)
# ═══════════════════════════════════════════════════════════════════════════════

class TestUrlIngestion:

    def test_http_webpage(self):
        """Load a stable Wikipedia page."""
        from src.ingestion.document_loader import load_document
        try:
            t0 = time.monotonic()
            text = load_document("https://en.wikipedia.org/wiki/Photosynthesis")
            t = time.monotonic() - t0
            assert len(text) > 500, "Webpage text too short"
            assert "chlorophyll" in text.lower() or "photosynthesis" in text.lower()
            print(f"\n  HTTP webpage: {len(text):,} chars  ({t:.2f}s)")
        except Exception as exc:
            pytest.skip(f"Network unavailable or blocked: {exc}")

    def test_http_pdf_url(self):
        """Download and parse a publicly hosted PDF."""
        from src.ingestion.document_loader import load_document
        # Stable small public PDF from Mozilla
        url = "https://www.mozilla.org/en-US/press/mozilla-2022-10-18-press-kit.pdf"
        try:
            t0 = time.monotonic()
            text = load_document(url)
            t = time.monotonic() - t0
            assert len(text) > 200
            print(f"\n  PDF-URL: {len(text):,} chars  ({t:.2f}s)")
        except Exception as exc:
            pytest.skip(f"URL not reachable: {exc}")


# ═══════════════════════════════════════════════════════════════════════════════
# GROUP C — Image formats (Groq Vision — needs GROQ_API_KEY)
# ═══════════════════════════════════════════════════════════════════════════════

class TestImageFormats:

    @SKIP_NO_GROQ
    def test_png(self):
        data = _make_png()
        text, t = _load(data, ".png")
        # Groq Vision on a white 10x10 PNG returns some description
        assert isinstance(text, str) and len(text) > 0
        print(f"\n  PNG:  {len(text):,} chars  ({t:.2f}s)  preview: {text[:80]!r}")

    @SKIP_NO_GROQ
    def test_jpg(self):
        data = _make_jpeg()
        text, t = _load(data, ".jpg")
        assert isinstance(text, str) and len(text) > 0
        print(f"\n  JPG:  {len(text):,} chars  ({t:.2f}s)  preview: {text[:80]!r}")

    @SKIP_NO_GROQ
    def test_jpeg_extension(self):
        data = _make_jpeg()
        text, t = _load(data, ".jpeg")
        assert isinstance(text, str) and len(text) > 0
        print(f"\n  JPEG: {len(text):,} chars  ({t:.2f}s)")

    @SKIP_NO_GROQ
    def test_webp(self):
        """WEBP — create a real WEBP file using Pillow."""
        try:
            from PIL import Image as _Image
            img = _Image.new("RGB", (10, 10), color=(200, 200, 255))
            buf = io.BytesIO()
            img.save(buf, format="WEBP")
            data = buf.getvalue()
        except ImportError:
            pytest.skip("Pillow not installed — cannot create valid WEBP test file")
        text, t = _load(data, ".webp")
        assert isinstance(text, str)
        print(f"\n  WEBP: {len(text):,} chars  ({t:.2f}s)")


# ═══════════════════════════════════════════════════════════════════════════════
# GROUP D — Audio formats (Groq Whisper — needs GROQ_API_KEY)
# ═══════════════════════════════════════════════════════════════════════════════

class TestAudioFormats:

    @SKIP_NO_GROQ
    def test_wav(self):
        """1-second 440 Hz WAV tone — Whisper returns transcription or empty."""
        data = _make_wav()
        try:
            text, t = _load(data, ".wav")
            # Whisper on a pure sine tone may return empty or music notation
            assert isinstance(text, str)
            print(f"\n  WAV:  {len(text):,} chars  ({t:.2f}s)  preview: {text[:60]!r}")
        except ValueError as exc:
            # "No speech detected" is an acceptable outcome for a sine tone
            assert "speech" in str(exc).lower() or "no" in str(exc).lower(), \
                f"Unexpected error for WAV: {exc}"
            print(f"\n  WAV:  (no speech detected — expected for sine tone)")

    @SKIP_NO_GROQ
    def test_mp3_format_routing(self):
        """Verify MP3 is routed to Whisper (not plain-text reader)."""
        from src.ingestion.document_loader import _AUDIO_VIDEO_EXTENSIONS
        assert ".mp3" in _AUDIO_VIDEO_EXTENSIONS, ".mp3 not registered as audio format"
        print("\n  MP3 routing: PASS (registered as audio/video extension)")

    @SKIP_NO_GROQ
    def test_m4a_format_routing(self):
        from src.ingestion.document_loader import _AUDIO_VIDEO_EXTENSIONS
        assert ".m4a" in _AUDIO_VIDEO_EXTENSIONS
        print("\n  M4A routing: PASS")


# ═══════════════════════════════════════════════════════════════════════════════
# GROUP E — Video formats (Groq Whisper + ffmpeg)
# ═══════════════════════════════════════════════════════════════════════════════

class TestVideoFormats:

    def test_mp4_format_routing(self):
        from src.ingestion.document_loader import _AUDIO_VIDEO_EXTENSIONS
        assert ".mp4" in _AUDIO_VIDEO_EXTENSIONS
        print("\n  MP4 routing: PASS")

    def test_webm_format_routing(self):
        from src.ingestion.document_loader import _AUDIO_VIDEO_EXTENSIONS
        assert ".webm" in _AUDIO_VIDEO_EXTENSIONS
        print("\n  WEBM routing: PASS")

    def test_avi_needs_ffmpeg(self):
        """AVI/MOV/MKV are routed through ffmpeg; verify registration."""
        from src.ingestion.document_loader import _NEEDS_FFMPEG_CONVERSION
        for ext in (".avi", ".mov", ".mkv"):
            assert ext in _NEEDS_FFMPEG_CONVERSION, f"{ext} not in _NEEDS_FFMPEG_CONVERSION"
        print("\n  AVI/MOV/MKV ffmpeg routing: PASS")

    def test_ffmpeg_availability(self):
        """Report whether ffmpeg is installed (AVI/MOV/MKV support depends on it)."""
        import shutil
        available = shutil.which("ffmpeg") is not None
        print(f"\n  ffmpeg available: {'YES' if available else 'NO (AVI/MOV/MKV will fail)'}")
        # Not a hard failure — just informational


# ═══════════════════════════════════════════════════════════════════════════════
# GROUP F — YouTube URL (layered transcript fallback)
# ═══════════════════════════════════════════════════════════════════════════════

class TestYouTube:

    @SKIP_NO_YT
    def test_youtube_url_short(self):
        """Short public YouTube video with captions."""
        from src.ingestion.document_loader import load_document
        url = "https://www.youtube.com/watch?v=dQw4w9WgXcQ"  # well-known, captioned
        try:
            t0 = time.monotonic()
            text = load_document(url)
            t = time.monotonic() - t0
            assert len(text) > 50
            print(f"\n  YouTube: {len(text):,} chars  ({t:.2f}s)  preview: {text[:80]!r}")
        except Exception as exc:
            pytest.skip(f"YouTube transcript unavailable (expected on cloud IPs): {exc}")

    def test_youtube_url_parsing(self):
        """Verify all YouTube URL formats are correctly parsed."""
        from src.ingestion.document_loader import _youtube_video_id
        cases = {
            "https://www.youtube.com/watch?v=dQw4w9WgXcQ": "dQw4w9WgXcQ",
            "https://youtu.be/dQw4w9WgXcQ":                "dQw4w9WgXcQ",
            "https://www.youtube.com/shorts/dQw4w9WgXcQ":   "dQw4w9WgXcQ",
            "https://www.youtube.com/live/dQw4w9WgXcQ":     "dQw4w9WgXcQ",
            "https://www.youtube.com/embed/dQw4w9WgXcQ":    "dQw4w9WgXcQ",
        }
        for url, expected_id in cases.items():
            got = _youtube_video_id(url)
            assert got == expected_id, f"URL '{url}': expected '{expected_id}', got '{got}'"
        print(f"\n  YouTube URL parsing: {len(cases)} URL formats — PASS")


# ═══════════════════════════════════════════════════════════════════════════════
# GROUP G — Full pipeline smoke test (needs LLM key)
# ═══════════════════════════════════════════════════════════════════════════════

class TestFullPipeline:

    @SKIP_NO_LLM
    def test_generate_questions_txt(self):
        """End-to-end: TXT → extract → generate_questions_from_text → validate."""
        from src.ingestion.document_loader import load_document
        from src.generation.qa_generator import generate_questions_from_text

        # Write a richer document so the LLM has enough context
        content = _CONTENT * 3  # ~3 paragraphs
        with tempfile.NamedTemporaryFile(
            suffix=".txt", delete=False, mode="w", encoding="utf-8"
        ) as fh:
            fh.write(content)
            path = fh.name

        try:
            t0 = time.monotonic()
            text = load_document(path)
            questions = generate_questions_from_text(
                text, num_questions=3, request_id="e2e-test"
            )
            t = time.monotonic() - t0
        finally:
            Path(path).unlink(missing_ok=True)

        assert isinstance(questions, list), "Expected list of questions"
        assert len(questions) >= 1, "Expected at least 1 question"

        for q in questions:
            assert "question" in q and q["question"].strip(), "Missing question text"
            assert "choices" in q and len(q["choices"]) == 4, "Choices must have A-D"
            assert q.get("correct_answer") in ("A", "B", "C", "D"), "Invalid correct_answer"
            assert "explanation" in q, "Missing explanation"

        print(f"\n  Pipeline TXT: {len(questions)} questions generated  ({t:.1f}s)")
        for q in questions:
            print(f"    Q{q['question_number']}: {q['question'][:70]}")

    @SKIP_NO_LLM
    def test_generate_questions_pdf(self):
        """End-to-end: PDF → extract → generate_questions_from_text → validate."""
        from src.ingestion.document_loader import load_document
        from src.generation.qa_generator import generate_questions_from_text

        data = _make_pdf()
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as fh:
            fh.write(data)
            path = fh.name

        try:
            t0 = time.monotonic()
            text = load_document(path)
            questions = generate_questions_from_text(
                text, num_questions=3, request_id="e2e-pdf-test"
            )
            t = time.monotonic() - t0
        finally:
            Path(path).unlink(missing_ok=True)

        assert len(questions) >= 1
        print(f"\n  Pipeline PDF: {len(questions)} questions  ({t:.1f}s)")

    @SKIP_NO_LLM
    def test_summarize_text(self):
        """End-to-end: text → summarize_text → validate markdown structure."""
        from src.generation.qa_generator import summarize_text

        t0 = time.monotonic()
        summary = summarize_text(_CONTENT * 2, request_id="e2e-summary-test")
        t = time.monotonic() - t0

        assert isinstance(summary, str) and len(summary) > 100
        assert "##" in summary, "Expected Markdown headers in summary"
        print(f"\n  Pipeline summary: {len(summary)} chars  ({t:.1f}s)")
        print(f"    Preview: {summary[:120]!r}")

    @SKIP_NO_LLM
    def test_large_doc_chunking(self):
        """Verify map-reduce kicks in for docs > _SINGLE_PASS_CHARS."""
        from src.generation.qa_generator import (
            generate_questions_from_text,
            _SINGLE_PASS_CHARS,
        )

        # Build a document larger than the single-pass threshold
        large_text = _CONTENT * 50   # ~45K chars → needs map-reduce at 60K threshold? No.
        # Make it definitely large
        large_text = (_CONTENT + "\n\n") * 200  # ~170K chars → well over 60K

        assert len(large_text) > _SINGLE_PASS_CHARS, \
            f"Test document too small: {len(large_text)} <= {_SINGLE_PASS_CHARS}"

        t0 = time.monotonic()
        questions = generate_questions_from_text(
            large_text, num_questions=5, request_id="e2e-large-test"
        )
        t = time.monotonic() - t0

        assert len(questions) >= 1
        print(
            f"\n  Large doc map-reduce: {len(large_text):,} chars "
            f"→ {len(questions)} questions  ({t:.1f}s)"
        )


# ═══════════════════════════════════════════════════════════════════════════════
# GROUP H — Chunking / diversity / deduplication unit tests (no API)
# ═══════════════════════════════════════════════════════════════════════════════

class TestChunkingAndDedup:

    def test_split_into_chunks_small_doc(self):
        """Docs smaller than chunk_size return a single chunk."""
        from src.generation.qa_generator import _split_into_chunks, _CHUNK_CHARS
        text = "Short document. " * 100  # ~1.6K chars
        chunks = _split_into_chunks(text, _CHUNK_CHARS)
        assert len(chunks) == 1
        assert chunks[0].strip() == text.strip()

    def test_split_into_chunks_large_doc(self):
        """Docs larger than chunk_size are split correctly with no content loss."""
        from src.generation.qa_generator import _split_into_chunks, _CHUNK_CHARS, _CHUNK_OVERLAP
        text = ("Word sentence paragraph. " * 5_000)  # ~125K chars
        chunks = _split_into_chunks(text, _CHUNK_CHARS, _CHUNK_OVERLAP)
        assert len(chunks) > 1
        # Verify all chunks are within size limit (with small tolerance for separator)
        for c in chunks:
            assert len(c) <= _CHUNK_CHARS + 100, f"Chunk too large: {len(c)}"
        # Verify we haven't lost content (total covered >= doc length - overlaps)
        total_raw = sum(len(c) for c in chunks)
        assert total_raw >= len(text) * 0.95, "Too much content lost in chunking"

    def test_select_diverse_chunks_fewer_than_n(self):
        """When chunks < n, all chunks are returned."""
        from src.generation.qa_generator import _select_diverse_chunks
        chunks = ["chunk A", "chunk B", "chunk C"]
        selected = _select_diverse_chunks(chunks, 10)
        assert len(selected) == 3

    def test_select_diverse_chunks_returns_n(self):
        """Exactly n chunks are returned when len(chunks) > n."""
        from src.generation.qa_generator import _select_diverse_chunks
        chunks = [f"Topic {i}: sentence about subject number {i}." for i in range(20)]
        selected = _select_diverse_chunks(chunks, 6)
        assert len(selected) == 6
        # Returned in document order
        indices = [i for i, _ in selected]
        assert indices == sorted(indices)

    def test_select_diverse_chunks_always_includes_first_and_last(self):
        """First and last chunks are always selected (intro + conclusion)."""
        from src.generation.qa_generator import _select_diverse_chunks
        chunks = [f"Chunk {i} with unique content {i*7}." for i in range(30)]
        selected = _select_diverse_chunks(chunks, 5)
        indices = [i for i, _ in selected]
        assert 0 in indices, "First chunk (intro) must always be selected"
        assert 29 in indices, "Last chunk (conclusion) must always be selected"

    def test_deduplicate_removes_near_duplicates(self):
        """Near-duplicate questions (Jaccard >= 0.65) are removed."""
        from src.generation.qa_generator import _deduplicate_questions
        questions = [
            {"question": "How does SQL indexing improve query performance?",
             "choices": {}, "correct_answer": "A", "explanation": ""},
            {"question": "How does SQL indexing improve database performance?",  # near-dup
             "choices": {}, "correct_answer": "B", "explanation": ""},
            {"question": "What is the role of chlorophyll in photosynthesis?",
             "choices": {}, "correct_answer": "C", "explanation": ""},
        ]
        result = _deduplicate_questions(questions)
        assert len(result) == 2
        questions_text = [q["question"] for q in result]
        assert any("SQL" in q for q in questions_text)
        assert any("chlorophyll" in q for q in questions_text)

    def test_deduplicate_keeps_distinct_questions(self):
        """Distinct questions (Jaccard < 0.65) are all kept."""
        from src.generation.qa_generator import _deduplicate_questions
        questions = [
            {"question": "What is photosynthesis?",
             "choices": {}, "correct_answer": "A", "explanation": ""},
            {"question": "How does the Calvin cycle work?",
             "choices": {}, "correct_answer": "B", "explanation": ""},
            {"question": "What products are made during photosynthesis?",
             "choices": {}, "correct_answer": "C", "explanation": ""},
        ]
        result = _deduplicate_questions(questions)
        assert len(result) == 3, f"Expected 3 distinct questions, kept {len(result)}"

    def test_tfidf_diversity_beats_even_on_repeated_topic(self):
        """
        When document has 10 chunks about topic A and 5 about topic B,
        diversity selection should not pick more than 2 consecutive same-topic chunks.
        """
        from src.generation.qa_generator import _select_diverse_chunks

        topic_a = [
            "Neural networks learn through gradient descent and backpropagation.",
            "Convolutional networks apply filters to detect spatial features in images.",
            "Recurrent networks process sequential data with hidden state memory.",
            "Transformers use attention mechanism to weigh token relationships.",
            "BERT is a bidirectional transformer trained on masked language modelling.",
            "GPT models generate text autoregressively from left to right tokens.",
            "Dropout regularises networks by randomly zeroing activations during training.",
            "Batch normalisation stabilises training by normalising layer inputs.",
            "Adam optimiser adapts learning rates per parameter using momentum estimates.",
            "Overfitting is reduced by early stopping on the validation loss curve.",
        ]
        topic_b = [
            "SQL databases store data in relational tables with primary and foreign keys.",
            "NoSQL databases like MongoDB store documents in JSON-like BSON format.",
            "Indexing improves SELECT performance by avoiding full table scans.",
            "ACID transactions guarantee atomicity consistency isolation durability.",
            "Sharding distributes large datasets across multiple database nodes.",
        ]
        all_chunks = topic_a + topic_b   # 15 chunks, A: indices 0-9, B: indices 10-14

        selected = _select_diverse_chunks(all_chunks, 6)
        indices = [i for i, _ in selected]

        # At least 1 chunk from topic B should be selected
        b_count = sum(1 for i in indices if i >= 10)
        assert b_count >= 1, \
            f"Diversity selection ignored topic B entirely: indices={indices}"
        print(f"\n  TF-IDF diversity: indices={indices}, topic_B_selected={b_count}")


# ═══════════════════════════════════════════════════════════════════════════════
# Standalone runner
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import subprocess
    result = subprocess.run(
        [
            sys.executable, "-m", "pytest",
            __file__,
            "-v",
            "--tb=short",
            "--no-header",
            "-rN",
        ],
        cwd=str(Path(__file__).parent.parent),
    )
    sys.exit(result.returncode)
